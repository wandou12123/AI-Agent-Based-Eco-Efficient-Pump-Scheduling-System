# -*- coding: utf-8 -*-
"""将项目详设 Markdown 转为 .docx（表格、标题、列表、代码块、加粗/行内代码）。"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


def set_doc_defaults(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)


def add_inline_runs(paragraph, text: str) -> None:
    if not text:
        return
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            try:
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            except Exception:
                pass
        else:
            r = paragraph.add_run(part)
            try:
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            except Exception:
                pass


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.count("|") >= 2


def is_table_sep(line: str) -> bool:
    s = line.strip().strip("|")
    if not s:
        return False
    return bool(re.match(r"^[\s|\-:]+$", line)) and "-" in line


def flush_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < cols:
            r.append("")
    t = document.add_table(rows=len(rows), cols=cols)
    t.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = t.rows[ri].cells[ci]
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            add_inline_runs(p, cell_text)
            for run in p.runs:
                run.font.size = Pt(9)


def parse_markdown_to_docx(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    document = Document()
    set_doc_defaults(document)
    document.sections[0].page_height = Cm(29.7)
    document.sections[0].page_width = Cm(21.0)
    document.sections[0].left_margin = Cm(2.54)
    document.sections[0].right_margin = Cm(2.54)
    document.sections[0].top_margin = Cm(2.54)
    document.sections[0].bottom_margin = Cm(2.54)

    i = 0
    in_code = False
    code_lines: list[str] = []

    def add_code_block(lines_block: list[str]) -> None:
        if not lines_block:
            return
        p = document.add_paragraph()
        run = p.add_run()
        for idx, code_line in enumerate(lines_block):
            run.add_text(code_line)
            if idx != len(lines_block) - 1:
                run.add_break()
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                add_code_block(code_lines)
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped == "---":
            p = document.add_paragraph()
            p.add_run("—" * 32).font.size = Pt(8)
            i += 1
            continue

        if is_table_row(line):
            rows: list[list[str]] = []
            while i < len(lines) and is_table_row(lines[i]):
                row_line = lines[i].strip()
                if is_table_sep(row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.split("|")]
                if cells and cells[0] == "":
                    cells = cells[1:]
                if cells and cells[-1] == "":
                    cells = cells[:-1]
                rows.append(cells)
                i += 1
            flush_table(document, rows)
            continue

        if stripped.startswith("#"):
            m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
            if m:
                level = len(m.group(1))
                title = m.group(2).strip()
                hl = {1: 0, 2: 1, 3: 2, 4: 3, 5: 3, 6: 3}.get(level, 3)
                if level == 1:
                    document.add_heading(title, level=0)
                elif level == 2:
                    document.add_heading(title, level=1)
                elif level == 3:
                    document.add_heading(title, level=2)
                else:
                    document.add_heading(title, level=3)
                i += 1
                continue

        if stripped.startswith(">"):
            content = stripped[1:].lstrip().strip()
            p = document.add_paragraph(style=None)
            p.paragraph_format.left_indent = Cm(0.8)
            add_inline_runs(p, content)
            i += 1
            continue

        if stripped.startswith("- "):
            content = stripped[2:].strip()
            p = document.add_paragraph(style="List Bullet")
            add_inline_runs(p, content)
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            content = re.sub(r"^\d+\.\s*", "", stripped)
            p = document.add_paragraph(style="List Number")
            add_inline_runs(p, content)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        p = document.add_paragraph()
        add_inline_runs(p, stripped)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))


def main() -> None:
    # 详设等报告在仓库外上一级 docs/ 目录
    repo_root = Path(__file__).resolve().parent.parent
    docs_dir = repo_root.parent / "docs"
    md = docs_dir / "智慧水利应用_详细设计报告_v1.1.md"
    if len(sys.argv) >= 2:
        md = Path(sys.argv[1])
    out = md.with_suffix(".docx")
    if len(sys.argv) >= 3:
        out = Path(sys.argv[2])
    parse_markdown_to_docx(md, out)
    print(f"Written: {out}")


if __name__ == "__main__":
    main()
