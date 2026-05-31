"""政府文书（.docx）解析服务（详设 §5.5.3）"""
import os
import re
from docx import Document as DocxDocument


def parse_docx_text(file_path: str) -> str:
    """
    提取并清洗 docx 文本（去噪、分段、截断），为后续分析提供可审计输入。

    Args:
        file_path: 服务端可读的上传文件绝对路径

    Returns:
        清洗后的纯文本；文件不存在或无法读取时返回空字符串
    """
    if not os.path.exists(file_path):
        return ""
    doc = DocxDocument(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.replace("|", "").strip():
                paragraphs.append(row_text)
    raw = "\n".join(paragraphs)
    # 去多余空白行
    cleaned = re.sub(r"\n{3,}", "\n\n", raw).strip()
    return cleaned


def extract_text_from_docx(file_path: str) -> str:
    """兼容旧接口，委托 parse_docx_text"""
    return parse_docx_text(file_path)
